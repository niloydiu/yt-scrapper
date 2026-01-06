import os
import json
import logging
import io
import csv
import re
from datetime import datetime
from collections import Counter

from flask import Flask, request, render_template_string, send_file, jsonify, make_response

# Import core modules
from scraper import get_video_metadata
from transcripts import fetch_transcript_and_translation
from parser import extract_ingredients
from utils import save_records_to_json

# Try importing python-docx
try:
    from docx import Document
except ImportError:
    Document = None

APP_DIR = os.path.dirname(__file__)
DATA_DIR = os.path.join(APP_DIR, "data")
os.makedirs(DATA_DIR, exist_ok=True)

# Configure Logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
log = logging.getLogger(__name__)

app = Flask(__name__)

# --- TEMPLATE ---
INDEX_HTML = '''
<!doctype html>
<html lang="en">
  <head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>Niloy — YT Scrapper</title>
    <link rel="icon" href="/static/logo.svg" type="image/svg+xml">
    <link href="https://cdn.jsdelivr.net/npm/bootstrap-icons@1.11.3/font/bootstrap-icons.css" rel="stylesheet">
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap" rel="stylesheet">
    <style>
      :root {
        --bg-body: #050505;
        --bg-card: #0F0F0F;
        --bg-input: #171717;
        --bg-hover: #1F1F1F;
        --border-color: #262626;
        --primary: #EDEDED;
        --primary-inv: #000000;
        --text-main: #EDEDED;
        --text-muted: #A1A1AA;
        --accent: #3b82f6; /* Blue accent */
        --success: #10b981;
        --danger: #ef4444;
      }

      * { margin: 0; padding: 0; box-sizing: border-box; }
      
      body {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
        background: var(--bg-body);
        color: var(--text-main);
        line-height: 1.5;
        padding-bottom: 120px;
        -webkit-font-smoothing: antialiased;
      }

      .container {
        max-width: 1400px;
        margin: 0 auto;
        padding: 0 24px;
      }

      /* TOAST NOTIFICATIONS */
      #toast-container {
        position: fixed; top: 24px; right: 24px; z-index: 9999;
        display: flex; flex-direction: column; gap: 12px;
      }
      .toast {
        background: rgba(23, 23, 23, 0.9);
        backdrop-filter: blur(12px);
        border: 1px solid var(--border-color);
        color: var(--text-main);
        padding: 16px 20px;
        border-radius: 12px;
        font-size: 14px;
        font-weight: 500;
        box-shadow: 0 10px 30px rgba(0,0,0,0.5);
        display: flex; align-items: center; gap: 12px;
        animation: slideIn 0.3s cubic-bezier(0.16, 1, 0.3, 1);
        min-width: 300px;
      }
      .toast.success { border-left: 4px solid var(--success); }
      .toast.error { border-left: 4px solid var(--danger); }
      .toast.info { border-left: 4px solid var(--accent); }
      
      @keyframes slideIn {
        from { opacity: 0; transform: translateX(20px); }
        to { opacity: 1; transform: translateX(0); }
      }
      @keyframes fadeOut {
        to { opacity: 0; transform: translateX(10px); }
      }

      /* HEADER & HERO */
      header { padding: 80px 0 40px; text-align: center; }
      .brand-badge {
        display: inline-flex; align-items: center; gap: 6px;
        padding: 6px 16px; background: rgba(255,255,255,0.05);
        border: 1px solid var(--border-color); border-radius: 100px;
        font-size: 12px; font-weight: 600; letter-spacing: 0.05em;
        color: var(--text-muted); margin-bottom: 24px;
      }
      h1 {
        font-size: 48px; font-weight: 700; letter-spacing: -0.02em;
        background: linear-gradient(180deg, #fff 0%, #a1a1aa 100%);
        -webkit-background-clip: text; -webkit-text-fill-color: transparent;
        margin-bottom: 16px;
      }
      .subtitle { color: var(--text-muted); font-size: 18px; max-width: 540px; margin: 0 auto; }

      /* SEARCH BAR AREA */
      .search-wrapper {
        max-width: 600px; margin: 48px auto; position: relative;
      }
      .search-box {
        background: var(--bg-card); border: 1px solid var(--border-color);
        border-radius: 16px; padding: 8px; display: flex; gap: 12px;
        box-shadow: 0 0 0 0 transparent; transition: all 0.2s;
      }
      .search-box:focus-within {
        border-color: #525252;
        box-shadow: 0 0 0 4px rgba(255,255,255,0.05);
      }
      .search-input {
        flex: 1; background: transparent; border: none; padding: 12px 16px;
        color: #fff; font-size: 16px; min-width: 0;
      }
      .search-input:focus { outline: none; }
      .btn-search {
        background: var(--primary); color: var(--primary-inv); border: none;
        border-radius: 10px; padding: 0 24px; font-weight: 600; cursor: pointer;
        transition: opacity 0.2s;
      }
      .btn-search:hover { opacity: 0.9; }
      .btn-search:disabled { opacity: 0.5; cursor: not-allowed; }

      /* STATUS BAR */
      .status-bar {
        max-width: 600px; margin: 0 auto 40px; display: none;
        background: #121212; border: 1px solid var(--border-color);
        border-radius: 12px; padding: 16px;
      }
      .progress-bg {
        height: 6px; background: #222; border-radius: 10px; overflow: hidden; margin-bottom: 8px;
      }
      .progress-fill {
        height: 100%; background: var(--success); width: 0%; transition: width 0.4s ease;
      }
      .status-label { font-size: 13px; color: var(--text-muted); text-align: center; }

      /* FILTERS & CONTROLS */
      .toolbar {
        display: none; align-items: center; justify-content: space-between;
        margin-bottom: 24px; padding-bottom: 24px; border-bottom: 1px solid var(--border-color);
      }
      .toolbar-left { display: flex; gap: 16px; align-items: center; }
      .toolbar-right { display: flex; gap: 12px; }
      
      .btn-text {
        background: transparent; border: 1px solid transparent; color: var(--text-muted);
        padding: 8px 12px; border-radius: 8px; cursor: pointer; font-size: 14px;
        transition: all 0.2s;
      }
      .btn-text:hover { color: #fff; background: rgba(255,255,255,0.05); }
      .btn-text.active { color: #fff; background: rgba(255,255,255,0.1); border-color: var(--border-color); }

      /* GRID */
      .grid {
        display: grid; grid-template-columns: repeat(auto-fill, minmax(280px, 1fr)); gap: 24px;
      }
      
      /* CARD STYLE */
      .card {
        background: var(--bg-card); border: 1px solid var(--border-color);
        border-radius: 16px; overflow: hidden; position: relative;
        transition: all 0.2s ease; cursor: default;
        display: flex; flex-direction: column;
      }
      .card:hover { transform: translateY(-2px); border-color: #404040; }
      .card.selected { border-color: var(--accent); box-shadow: 0 0 0 1px var(--accent); }

      .card-thumb {
        position: relative; aspect-ratio: 16/9; background: #111; cursor: pointer;
      }
      .card-thumb img { width: 100%; height: 100%; object-fit: cover; }
      
      .duration-badge {
        position: absolute; bottom: 8px; right: 8px;
        background: rgba(0,0,0,0.75); backdrop-filter: blur(4px);
        padding: 2px 6px; border-radius: 4px; font-size: 11px; font-weight: 600;
      }
      
      .select-overlay {
        position: absolute; inset: 0; background: rgba(0,0,0,0.3);
        display: none;
      }

      /* Visible top-left checkbox to make selection affordance explicit */
      .select-checkbox {
        position: absolute; top: 10px; left: 10px; z-index: 20;
        width: 28px; height: 28px; border-radius: 6px; display: flex; align-items: center; justify-content: center;
        background: rgba(0,0,0,0.5); border: 2px solid rgba(255,255,255,0.2); cursor: pointer;
      }
      .select-checkbox i { color: transparent; font-size: 14px; }
      .card.selected .select-checkbox { background: var(--accent); border-color: var(--accent); }
      .card.selected .select-checkbox i { color: #fff; }

      .card-body { padding: 16px; flex: 1; display: flex; flex-direction: column; }
      .card-title {
        font-size: 14px; font-weight: 600; line-height: 1.4; margin-bottom: 8px;
        display: -webkit-box; -webkit-line-clamp: 2; -webkit-box-orient: vertical; overflow: hidden;
      }
      .card-meta { margin-top: auto; display: flex; gap: 6px; flex-wrap: wrap; }
      .badge {
        font-size: 10px; padding: 2px 8px; border-radius: 4px; background: #27272a; color: #a1a1aa;
      }
      .badge.green { background: rgba(16, 185, 129, 0.15); color: #10b981; }
      
      /* BOTTOM ACTION BAR (Floating) */
      .fab-bar {
        position: fixed; bottom: 32px; left: 50%; transform: translateX(-50%) translateY(200%);
        background: rgba(23, 23, 23, 0.85); backdrop-filter: blur(20px);
        border: 1px solid rgba(255,255,255,0.1); padding: 8px 8px 8px 24px;
        border-radius: 100px; display: flex; align-items: center; gap: 24px;
        box-shadow: 0 24px 48px -12px rgba(0,0,0,0.8); z-index: 100;
        transition: transform 0.4s cubic-bezier(0.19, 1, 0.22, 1);
        min-width: 400px; justify-content: space-between;
      }
      .fab-bar.visible { transform: translateX(-50%) translateY(0); }
      
      .selection-info { font-size: 14px; font-weight: 600; color: #fff; display: flex; gap: 8px; align-items: center; }
      .selection-actions { display: flex; gap: 8px; }
      
      .btn-fab {
        height: 40px; padding: 0 20px; border-radius: 100px; border: none;
        font-size: 13px; font-weight: 600; cursor: pointer; display: flex; align-items: center; gap: 8px;
        transition: all 0.2s; white-space: nowrap;
      }
      .btn-fab.secondary { background: rgba(255,255,255,0.08); color: #fff; }
      .btn-fab.secondary:hover { background: rgba(255,255,255,0.12); }
      .btn-fab.primary { background: #fff; color: #000; }
      .btn-fab.primary:hover { background: #e5e5e5; }
      
      /* MODAL SYSTEM */
      .modal-backdrop {
        position: fixed; inset: 0; background: rgba(0,0,0,0.8); backdrop-filter: blur(8px);
        z-index: 200; opacity: 0; pointer-events: none; transition: opacity 0.3s;
        display: flex; align-items: center; justify-content: center; padding: 20px;
      }
      .modal-backdrop.active { opacity: 1; pointer-events: auto; }
      
      .modal-window {
        background: var(--bg-card); border: 1px solid var(--border-color);
        width: 100%; max-width: 500px; border-radius: 24px; padding: 32px;
        transform: scale(0.95); transition: transform 0.3s cubic-bezier(0.19, 1, 0.22, 1);
        max-height: 90vh; display: flex; flex-direction: column;
      }
      .modal-backdrop.active .modal-window { transform: scale(1); }
      
      .modal-header { display: flex; justify-content: space-between; align-items: center; margin-bottom: 24px; }
      .modal-title { font-size: 20px; font-weight: 700; }
      .btn-close { background: none; border: none; color: var(--text-muted); cursor: pointer; font-size: 20px; }
      .btn-close:hover { color: #fff; }
      
      .modal-content-scroll { overflow-y: auto; padding-right: 8px; margin: -4px -8px -4px -4px; padding: 4px 8px 4px 4px; }
      
      /* FORM ELEMENTS */
      .form-group { margin-bottom: 24px; }
      .label-sm { font-size: 11px; font-weight: 700; text-transform: uppercase; color: var(--text-muted); letter-spacing: 0.05em; margin-bottom: 12px; display: block; }
      
      .check-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 12px; }
      .check-item {
        display: flex; align-items: center; gap: 10px; cursor: pointer;
        padding: 10px; border-radius: 8px; background: rgba(255,255,255,0.03);
        transition: background 0.2s;
      }
      .check-item:hover { background: rgba(255,255,255,0.06); }
      .check-item input { accent-color: #fff; width: 16px; height: 16px; }
      
      .select-box {
        width: 100%; background: var(--bg-input); border: 1px solid var(--border-color);
        padding: 12px; border-radius: 8px; color: #fff; font-size: 14px; outline: none;
      }
      
      .shop-list-item {
        display: flex; justify-content: space-between; padding: 8px 0;
        border-bottom: 1px solid #222; font-size: 14px;
      }
      .shop-list-item:last-child { border-bottom: none; }
      
      .actions-mt { margin-top: 32px; display: flex; gap: 12px; justify-content: flex-end; }
      .hero-meta { font-size: 16px; color: var(--text-muted); margin-top: 12px; }
      
      /* FOOTER */
      footer {
        margin-top: 160px;
        padding: 80px 0 64px;
        border-top: 1px solid var(--border-color);
        background: linear-gradient(180deg, transparent 0%, rgba(15,15,15,0.4) 100%);
      }
      .footer-container {
        max-width: 900px;
        margin: 0 auto;
        text-align: center;
      }
      .footer-brand {
        margin-bottom: 48px;
      }
      .footer-brand h2 {
        font-size: 28px;
        font-weight: 700;
        letter-spacing: -0.03em;
        margin-bottom: 16px;
        color: #fff;
        background: linear-gradient(135deg, #ffffff 0%, #a1a1aa 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
      }
      .footer-brand p {
        color: var(--text-muted);
        font-size: 15px;
        line-height: 1.7;
        max-width: 520px;
        margin: 0 auto;
      }
      .footer-divider {
        height: 1px;
        background: linear-gradient(90deg, transparent 0%, var(--border-color) 50%, transparent 100%);
        margin: 48px 0;
      }
      .footer-links {
        display: flex;
        justify-content: center;
        gap: 48px;
        flex-wrap: wrap;
        margin-bottom: 48px;
      }
      .footer-link-item {
        display: flex;
        flex-direction: column;
        align-items: center;
        gap: 8px;
        text-decoration: none;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
        padding: 20px 32px;
        border-radius: 12px;
        background: rgba(255,255,255,0.02);
        border: 1px solid transparent;
      }
      .footer-link-item:hover {
        background: rgba(255,255,255,0.05);
        border-color: rgba(255,255,255,0.1);
        transform: translateY(-2px);
      }
      .footer-link-icon {
        width: 44px;
        height: 44px;
        border-radius: 10px;
        background: rgba(255,255,255,0.06);
        display: flex;
        align-items: center;
        justify-content: center;
        transition: all 0.3s;
      }
      .footer-link-item:hover .footer-link-icon {
        background: rgba(59,130,246,0.15);
        transform: scale(1.1);
      }
      .footer-link-icon i {
        font-size: 20px;
        color: var(--text-muted);
        transition: color 0.3s;
      }
      .footer-link-item:hover .footer-link-icon i {
        color: var(--accent);
      }
      .footer-link-label {
        font-size: 11px;
        font-weight: 600;
        text-transform: uppercase;
        letter-spacing: 0.1em;
        color: var(--text-muted);
        transition: color 0.3s;
      }
      .footer-link-value {
        font-size: 14px;
        font-weight: 600;
        color: #fff;
        transition: color 0.3s;
      }
      .footer-link-item:hover .footer-link-label,
      .footer-link-item:hover .footer-link-value {
        color: #fff;
      }
      @media (max-width: 768px) {
        footer { margin-top: 120px; padding: 60px 0 48px; }
        .footer-brand h2 { font-size: 24px; }
        .footer-links { flex-direction: column; gap: 16px; }
        .footer-link-item { padding: 16px 24px; width: 100%; }
      }
    </style>
  </head>
  <body>
    <!-- TOAST CONTAINER -->
    <div id="toast-container"></div>

    <div class="container">
      <header>
        <div class="brand-badge"><img src="/static/logo.svg" alt="logo" style="height:20px; border-radius:6px; margin-right:8px;"> Niloy Kumar Mohonta</div>
        <h1>Capture YouTube metadata with confidence.</h1>
        <p class="subtitle">Niloy has made it easy to turn any YouTube channel into structured metadata, narratives, and extracted highlights.</p>
        <p class="hero-meta">
          Reach Niloy at <a href="mailto:niloykumarmohonta@gmail.com" style="color:#f3f4f6; text-decoration:underline;">niloykumarmohonta@gmail.com</a> — <a href="https://niloykm.vercel.app" style="color:#f3f4f6; text-decoration:underline;" target="_blank" rel="noreferrer">niloykm.vercel.app</a>
        </p>
      
        <div class="search-wrapper">
          <div class="search-box">
            <input id="urlInput" class="search-input" type="text" placeholder="Paste Channel or Playlist URL...">
            <button id="searchBtn" class="btn-search" onclick="fetchMetadata()">Search</button>
          </div>
            <div class="flow-steps" style="margin-top:12px; text-align:center; color:var(--text-muted); font-size:13px;">
            <strong>Quick Steps:</strong> 1) Search &nbsp;→&nbsp; 2) Select videos &nbsp;→&nbsp; 3) Click "Fetch Info" to retrieve transcripts &amp; extracted items &nbsp;→&nbsp; 4) Export or open Aggregated Items
          </div>
        </div>

        <div id="statusBar" class="status-bar">
          <div class="progress-bg"><div id="progressFill" class="progress-fill"></div></div>
          <div id="statusText" class="status-label">Initializing...</div>
        </div>
      </header>

      <div id="toolbar" class="toolbar">
        <div class="toolbar-left">
          <span style="font-size:14px; font-weight:600; color:#fff" id="resultCount">0 Videos</span>
          <button class="btn-text" onclick="toggleSelectAll()">Select All</button>
        </div>
        <div class="toolbar-right">
          <!-- Future: Sort/Filter -->
        </div>
      </div>

      <div id="grid" class="grid"></div>
    </div>

    <footer>
      <div class="container">
        <div class="footer-container">
          <div class="footer-brand">
            <h2>Niloy Kumar Mohonta</h2>
            <p>YouTube metadata, transcripts, and extracted items — curated by Niloy.</p>
          </div>
          
          <div class="footer-divider"></div>
          
          <div class="footer-links">
            <a href="mailto:niloykumarmohonta@gmail.com" class="footer-link-item">
              <div class="footer-link-icon">
                <i class="bi bi-envelope-at-fill"></i>
              </div>
              <span class="footer-link-label">Email</span>
              <span class="footer-link-value">niloykumarmohonta@gmail.com</span>
            </a>
            
            <a href="https://niloykm.vercel.app" target="_blank" rel="noreferrer" class="footer-link-item">
              <div class="footer-link-icon">
                <i class="bi bi-globe2"></i>
              </div>
              <span class="footer-link-label">Portfolio</span>
              <span class="footer-link-value">niloykm.vercel.app</span>
            </a>
            
            <a href="https://github.com/niloydiu" target="_blank" rel="noreferrer" class="footer-link-item">
              <div class="footer-link-icon">
                <i class="bi bi-github"></i>
              </div>
              <span class="footer-link-label">GitHub</span>
              <span class="footer-link-value">github.com/niloydiu</span>
            </a>
          </div>
        </div>
      </div>
    </footer>

    <!-- FLOATING ACTION BAR -->
    <div id="fabBar" class="fab-bar">
      <div class="selection-info">
        <i class="bi bi-check-circle-fill" style="color:var(--success)"></i>
        <span id="selectCount">0 selected</span>
      </div>
      
      <div class="selection-actions">
        <!-- Get Ingredients -->
        <button id="fetchInfoBtn" class="btn-fab secondary" onclick="handleFetchDetails()" title="Fetch transcripts and extract ingredients for selected videos">
          <i class="bi bi-magic"></i> Fetch Info
        </button>

        <!-- Aggregated Items -->
        <button id="shoppingBtn" class="btn-fab secondary" onclick="openShoppingList()" title="Open aggregated items from selected videos">
          <i class="bi bi-basket"></i> Aggregated Items
        </button>
        
        <!-- Export -->
        <button id="exportBtn" class="btn-fab primary" onclick="openExportModal()" title="Export selected videos in chosen format">
          <i class="bi bi-download"></i> Export
        </button>
      </div>
    </div>

    <!-- EXPORT MODAL -->
    <div id="exportModal" class="modal-backdrop" onclick="closeModal(event, 'exportModal')">
      <div class="modal-window">
        <div class="modal-header">
          <h2 class="modal-title">Export Data</h2>
          <button class="btn-close" onclick="closeModal(null, 'exportModal')"><i class="bi bi-x"></i></button>
        </div>
        
        <div class="form-group">
          <span class="label-sm">File Format</span>
          <select id="exportFormat" class="select-box">
            <option value="json">JSON (Data Object)</option>
            <option value="txt">Plain Text (.txt)</option>
            <option value="csv">CSV Spreadsheet (.csv)</option>
            <option value="docx">Microsoft Word (.docx)</option>
          </select>
        </div>

        <div class="form-group">
          <span class="label-sm">Include Fields</span>
          <div class="check-grid">
            <label class="check-item"><input type="checkbox" checked value="title"> Title</label>
            <label class="check-item"><input type="checkbox" checked value="link"> Link</label>
            <label class="check-item"><input type="checkbox" value="id"> Video ID</label>
            <label class="check-item"><input type="checkbox" value="duration"> Duration</label>
            <label class="check-item"><input type="checkbox" checked value="ingredients"> Items</label>
            <label class="check-item"><input type="checkbox" checked value="transcript"> Transcript</label>
          </div>
        </div>

        <div class="actions-mt">
           <button class="btn-fab secondary" onclick="closeModal(null, 'exportModal')">Cancel</button>
           <button class="btn-fab secondary" onclick="performExport('copy', this)">Copy Text</button>
           <button class="btn-fab primary" onclick="performExport('download', this)">Download</button>
        </div>
      </div>
    </div>

    <!-- SHOPPING LIST MODAL -->
    <div id="shopModal" class="modal-backdrop" onclick="closeModal(event, 'shopModal')">
      <div class="modal-window">
        <div class="modal-header">
          <h2 class="modal-title">Aggregated Items</h2>
          <button class="btn-close" onclick="closeModal(null, 'shopModal')"><i class="bi bi-x"></i></button>
        </div>
        <p style="color:#a1a1aa; font-size:13px; margin-bottom:16px;">
          Aggregated extracted items from selected videos. Note: These are extracted using heuristics and may contain non-relevant words.
        </p>
        <div id="shopListContent" class="modal-content-scroll" style="max-height: 400px;">
          <!-- Items go here -->
        </div>
        <div class="actions-mt">
          <button class="btn-fab secondary" onclick="copyShoppingList()">Copy List</button>
        </div>
      </div>
    </div>

    <script>
      // --- STATE ---
      let allVideos = [];
      let selectedIds = new Set();
      
      // --- DOM ---
      const grid = document.getElementById('grid');
      const fabBar = document.getElementById('fabBar');
      const statusBar = document.getElementById('statusBar');
      const progressFill = document.getElementById('progressFill');
      const statusText = document.getElementById('statusText');
      
      // --- TOAST SYSTEM ---
      function toast(msg, type='info') {
        const c = document.getElementById('toast-container');
        const t = document.createElement('div');
        t.className = `toast ${type}`;
        
        let icon = 'bi-info-circle-fill';
        if (type === 'success') icon = 'bi-check-circle-fill';
        if (type === 'error') icon = 'bi-exclamation-triangle-fill';

        t.innerHTML = `<i class="bi ${icon}"></i> <span>${msg}</span>`;
        c.appendChild(t);
        
        // Remove after 3s
        setTimeout(() => {
            t.style.animation = 'fadeOut 0.3s forwards';
            setTimeout(() => t.remove(), 300);
        }, 3000);
      }

      // --- LOGIC ---

      async function fetchMetadata() {
        const url = document.getElementById('urlInput').value.trim();
        if (!url) return toast('Please enter a valid YouTube URL', 'error');

        // Reset
        selectedIds.clear();
        updateFab();
        grid.innerHTML = '';
        
        setLoading(true, 'Extracting video metadata...');
        
        try {
          const res = await fetch('/api/extract', {
            method: 'POST',
            headers: { 'Content-Type': 'application/json' },
            body: JSON.stringify({ channel: url })
          });
          const data = await res.json();
          if (data.error) throw new Error(data.error);
          
          allVideos = data.results;
          if (allVideos.length === 0) toast('No videos found on this channel', 'info');
          else toast(`Found ${allVideos.length} videos`, 'success');

          document.getElementById('toolbar').style.display = 'flex';
          document.getElementById('resultCount').textContent = `${allVideos.length} Videos`;
          
          renderGrid();

        } catch (e) {
          toast(e.message, 'error');
        } finally {
          setLoading(false);
        }
      }

      async function handleFetchDetails() {
        if (selectedIds.size === 0) return toast('Select videos first', 'error');
        
        const videosToProcess = allVideos.filter(v => selectedIds.has(v.videoId));
        
        // Optimistic UI for unblocking
        setLoading(true, `Analysing ${videosToProcess.length} videos. This can take a minute...`);
        
        try {
           const res = await fetch('/api/fetch_details', {
            method: 'POST',
            headers: {'Content-Type': 'application/json'},
            body: JSON.stringify({ videos: videosToProcess })
           });
           const data = await res.json();
           
           // Update local state
           data.results.forEach(updated => {
             const idx = allVideos.findIndex(v => v.videoId === updated.videoId);
             if (idx !== -1) allVideos[idx] = updated;
           });

           renderGrid(); // Re-render to show badges
           toast('Transcripts & Ingredients fetched!', 'success');
        } catch(e) {
          toast(e.message, 'error');
        } finally {
          setLoading(false);
        }
      }

      function openShoppingList() {
        // Collect ingredients
        const ingredients = [];
        allVideos.filter(v => selectedIds.has(v.videoId)).forEach(v => {
            if (v.ingredients && Array.isArray(v.ingredients)) {
                ingredients.push(...v.ingredients);
            }
        });

        if (ingredients.length === 0) {
            toast('No ingredients found yet. Use "Fetch Info" first!', 'error');
            return;
        }

        // Simple aggregation (dedupe case insensitive)
        const counts = {};
        ingredients.forEach(i => {
           const norm = i.toLowerCase().trim();
           counts[norm] = (counts[norm] || 0) + 1;
        });
        
        const listContainer = document.getElementById('shopListContent');
        listContainer.innerHTML = '';
        
        // Sort alphabetically
        Object.keys(counts).sort().forEach(ing => {
            const div = document.createElement('div');
            div.className = 'shop-list-item';
            div.innerHTML = `<span>${ing}</span> <span style="color:#555">x${counts[ing]}</span>`;
            listContainer.appendChild(div);
        });

        // Open Modal
        document.getElementById('shopModal').classList.add('active');
      }
      
      function copyShoppingList() {
         const items = document.querySelectorAll('.shop-list-item span:first-child');
         const text = Array.from(items).map(i => i.textContent).join('\\n');
         navigator.clipboard.writeText(text);
         toast('Shopping list copied', 'success');
      }

      async function performExport(action, btnEl) {
        if (selectedIds.size === 0) return toast('No videos selected', 'error');

        const format = document.getElementById('exportFormat').value;
        const fields = Array.from(document.querySelectorAll('.check-grid input:checked')).map(cb => cb.value);
        
        if (fields.length === 0) return toast('Please select at least one field', 'error');

        const videosToExport = allVideos.filter(v => selectedIds.has(v.videoId));
        
        // Show immediate feedback
        const btn = btnEl || document.activeElement || document.querySelector('.btn-fab.primary');
        const originalText = btn && btn.innerText ? btn.innerText : '';
        if (btn) { btn.disabled = true; btn.innerText = 'Processing...'; }

        try {
            if (action === 'copy') {
                // Determine format -> force txt if copy? usually yes because you cant copy valid json/binary well to clipboard
                // But user might want JSON text
                let reqFormat = format;
                if (format === 'docx' || format === 'csv') reqFormat = 'txt'; // Fallback for copy

                const res = await fetch('/api/export', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({ videos: videosToExport, format: reqFormat, fields })
                });
                
                if (reqFormat === 'json') {
                    const json = await res.json();
                    navigator.clipboard.writeText(JSON.stringify(json, null, 2));
                } else {
                    const text = await res.text();
                    navigator.clipboard.writeText(text);
                }
                toast('Copied to Clipboard!', 'success');
            } else {
                // DOWNLOAD
                const res = await fetch('/api/export', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({ videos: videosToExport, format, fields })
                });
                
                if (!res.ok) throw new Error('Export failed on server');

                const blob = await res.blob();
                const url = window.URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url;
                
                // Extension
                let ext = format;
                // filename
                a.download = `export_${Date.now()}.${ext}`;
                document.body.appendChild(a);
                a.click();
                a.remove();
                toast('Download started', 'success');
            }
            closeModal(null, 'exportModal');
        } catch (e) {
            toast('Export failed: ' + e.message, 'error');
        } finally {
              if (btn) { btn.disabled = false; btn.innerText = originalText; }
        }
      }

      // --- RENDERING ---
      function renderGrid() {
        grid.innerHTML = '';
        allVideos.forEach(v => {
          const isSel = selectedIds.has(v.videoId);
          const hasIng = v.ingredients && v.ingredients.length > 0;
          const hasTx = !!v.transcript;
          
            const el = document.createElement('div');
            el.className = `card fade-in ${isSel ? 'selected' : ''}`;
          
            // Only toggle on click but avoid toggling when clicking links in future
            el.addEventListener('click', (e) => toggleSelection(v.videoId, e));

            el.innerHTML = `
            <div class="card-thumb">
              <div class="select-checkbox" title="Select video">
                <i class="bi bi-check"></i>
              </div>
              <img src="${v.thumbnail || ''}" loading="lazy">
              <div class="duration-badge">${v.duration || '0:00'}</div>
            </div>
            <div class="card-body">
              <div class="card-title">${v.title}</div>
              <div class="card-meta">
                 ${hasIng ? '<span class="badge green">Ingredients</span>' : ''}
                 ${hasTx ? '<span class="badge">Transcribed</span>' : ''}
              </div>
            </div>
            `;
          grid.appendChild(el);
        });
      }

      function toggleSelection(id, e) {
        if (selectedIds.has(id)) selectedIds.delete(id);
        else selectedIds.add(id);
        renderGrid();
        updateFab();
      }
      
      function toggleSelectAll() {
        if (selectedIds.size === allVideos.length) selectedIds.clear();
        else allVideos.forEach(v => selectedIds.add(v.videoId));
        renderGrid();
        updateFab();
      }
      
      function updateFab() {
        const c = selectedIds.size;
        document.getElementById('selectCount').innerText = `${c} selected`;
        if (c > 0) fabBar.classList.add('visible');
        else fabBar.classList.remove('visible');
        // Enable/disable action buttons based on selection
        const fetchBtn = document.getElementById('fetchInfoBtn');
        const shopBtn = document.getElementById('shoppingBtn');
        const expBtn = document.getElementById('exportBtn');
        if (fetchBtn) fetchBtn.disabled = c === 0;
        if (shopBtn) shopBtn.disabled = c === 0;
        if (expBtn) expBtn.disabled = c === 0;
        if (fetchBtn) fetchBtn.title = c === 0 ? 'Select videos first' : 'Fetch transcripts and ingredients';
        if (shopBtn) shopBtn.title = c === 0 ? 'Select videos first' : 'Open aggregated shopping list';
        if (expBtn) expBtn.title = c === 0 ? 'Select videos first' : 'Export selected videos';
      }

      function setLoading(bool, txt) {
        statusBar.style.display = bool ? 'block' : 'none';
        if (bool) {
            statusText.textContent = txt;
            progressFill.style.width = '30%';
            setTimeout(() => progressFill.style.width = '80%', 500);
        } else {
             progressFill.style.width = '100%';
        }
      }
      
      // Allow pressing Enter in the URL input to start the search
      (function attachEnterKey() {
        const input = document.getElementById('urlInput');
        if (!input) return;
        input.addEventListener('keydown', function (e) {
          if (e.key === 'Enter') {
            e.preventDefault();
            // small debounce to avoid double clicks
            if (!document.getElementById('searchBtn').disabled) fetchMetadata();
          }
        });
      })();

      // Modals
        function openExportModal() { 
          if(selectedIds.size===0) return toast('Select videos first','error');
          // ensure export button state matches selection
          document.getElementById('exportModal').classList.add('active'); 
        }
      function closeModal(e, id) {
          if (!e || e.target.classList.contains('modal-backdrop') || e.target.closest('.btn-close') || e.target.classList.contains('btn-fab')) {
             document.getElementById(id).classList.remove('active');
          }
      }
    </script>
  </body>
</html>
'''

# --- HELPERS ---

def format_text_output(videos, fields):
    out = []
    for v in videos:
        parts = []
        if 'title' in fields: parts.append(f"TITLE: {v.get('title')}")
        if 'link' in fields: parts.append(f"LINK: {v.get('youtubeVideoLink')}")
        if 'id' in fields: parts.append(f"ID: {v.get('videoId')}")
        if 'duration' in fields: parts.append(f"DURATION: {v.get('duration', '')}")
        
        if 'ingredients' in fields and v.get('ingredients'):
            parts.append("INGREDIENTS:")
            for i in v.get('ingredients', []):
                parts.append(f"- {i}")
        
        if 'transcript' in fields and v.get('transcript'):
            parts.append("TRANSCRIPT (Excerpt):")
            # Only 500 chars to keep txt readable, or full if user wants?
            # User probably wants full if they export to txt.
            parts.append(v.get('transcript', ''))
            
        out.append("\\n".join(parts))
        out.append("-" * 50)
    return "\\n".join(out)

def create_docx(videos, fields):
    if not Document: return None
    doc = Document()
    doc.add_heading('YouTube Export', 0)
    doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    for v in videos:
        if 'title' in fields: 
            doc.add_heading(v.get('title', 'Untitled'), level=1)
        
        # Meta info paragraph
        meta = []
        if 'link' in fields: meta.append(f"Link: {v.get('youtubeVideoLink')}")
        if 'duration' in fields: meta.append(f"Duration: {v.get('duration')}")
        if meta:
            p = doc.add_paragraph()
            for m in meta: p.add_run(m + "\\n")
            
        if 'ingredients' in fields and v.get('ingredients'):
            doc.add_heading('Ingredients', level=2)
            for ing in v.get('ingredients', []):
                doc.add_paragraph(ing, style='List Bullet')
                
        if 'transcript' in fields and v.get('transcript'):
            doc.add_heading('Transcript', level=2)
            doc.add_paragraph(v.get('transcript', ''))
            
        doc.add_paragraph("_" * 20)
        
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f

# --- ROUTES ---

@app.route('/')
def index():
    return render_template_string(INDEX_HTML)

@app.route('/api/extract', methods=['POST'])
def api_extract():
    data = request.get_json(silent=True) or {}
    channel = data.get('channel')
    if not channel: return jsonify({'error': 'Missing URL'}), 400

    try:
        videos = get_video_metadata(channel)
        return jsonify({'results': videos})
    except Exception as e:
        log.error(f"Extract error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/fetch_details', methods=['POST'])
def api_fetch_details():
    try:
        data = request.get_json(silent=True) or {}
        videos = data.get('videos', [])
        
        updated = []
        for vid in videos:
            # Reconstruct transcript/ingredients if missing but heuristic data exists? 
            # No, assume client sends what it has.
            
            if vid.get('ingredients'): # Skip if already has data
                updated.append(vid)
                continue
                
            vid_id = vid.get('videoId')
            try:
                # Fetch
                transcript, translation = fetch_transcript_and_translation(vid_id)
                text = translation if translation else transcript
                # Parse
                ingredients = extract_ingredients(text)
                
                vid['transcript'] = transcript
                vid['translation'] = translation
                vid['ingredients'] = ingredients
                updated.append(vid)
            except Exception as e:
                log.warning(f"Failed details for {vid_id}: {e}")
                vid['error'] = str(e)
                updated.append(vid)
                
        return jsonify({'results': updated})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/export', methods=['POST'])
def api_export():
    try:
        data = request.get_json(silent=True) or {}
        videos = data.get('videos', [])
        fmt = data.get('format', 'json')
        fields = data.get('fields', [])
        
        log.info(f"Exporting {len(videos)} videos to {fmt}")

        if fmt == 'json':
            # Clean selection based on fields
            clean_list = []
            for v in videos:
                item = {}
                # Map fields to keys
                if 'title' in fields: item['title'] = v.get('title')
                if 'link' in fields: item['youtubeVideoLink'] = v.get('youtubeVideoLink')
                if 'id' in fields: item['videoId'] = v.get('videoId')
                if 'duration' in fields: item['duration'] = v.get('duration')
                if 'ingredients' in fields: item['ingredients'] = v.get('ingredients')
                if 'transcript' in fields: item['transcript'] = v.get('transcript')
                clean_list.append(item)
            return jsonify(clean_list)

        elif fmt == 'txt':
            content = format_text_output(videos, fields)
            resp = make_response(content)
            resp.headers['Content-Type'] = 'text/plain; charset=utf-8'
            return resp

        elif fmt == 'csv':
            si = io.StringIO()
            writer = csv.writer(si)
            # Dynamic header
            header = [f.upper() for f in fields]
            writer.writerow(header)
            
            for v in videos:
                row = []
                for f in fields:
                    key = 'youtubeVideoLink' if f == 'link' else 'videoId' if f == 'id' else f
                    val = v.get(key, '')
                    if isinstance(val, list): val = "; ".join(val) # CSV friendly
                    row.append(val)
                writer.writerow(row)
                
            resp = make_response(si.getvalue())
            resp.headers['Content-Type'] = 'text/csv; charset=utf-8'
            return resp
            
        elif fmt == 'docx':
            if not Document: return jsonify({'error': 'python-docx not installed'}), 500
            file_stream = create_docx(videos, fields)
            return send_file(
                file_stream,
                as_attachment=True,
                download_name='export.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
        return jsonify({'error': 'Invalid format'}), 400

    except Exception as e:
        log.error(f"Export failure: {e}")
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(host='127.0.0.1', port=8000, debug=True)
